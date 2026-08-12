import asyncio
import ipaddress
import os
import json
import socket
import uuid
from typing import List
from urllib.parse import urljoin, urlparse
from urllib.robotparser import RobotFileParser
from xml.etree import ElementTree
import httpx
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException
from ..models.document import Document, DocumentChunk
from ..core.config import settings
from ..core.database import SessionLocal
from ..rag.embeddings import embed_texts
from . import plan_service


ALLOWED_TYPES = {"pdf", "docx", "txt", "csv", "xlsx", "url"}
MAX_URL_FETCH_BYTES = 5 * 1024 * 1024

# Hard operational ceiling on a single "import my whole website" crawl,
# independent of the business's plan (which caps total documents overall,
# but could be unlimited on higher tiers -- this still bounds how much work
# one crawl request can trigger).
MAX_CRAWL_PAGES = 40
CRAWL_USER_AGENT = "MielikkiXBot/1.0"
_NON_PAGE_EXTENSIONS = (
    ".pdf", ".jpg", ".jpeg", ".png", ".gif", ".svg", ".webp", ".ico",
    ".zip", ".mp4", ".mp3", ".css", ".js", ".xml", ".json", ".woff", ".woff2",
)


def _extract_text(path: str, file_type: str) -> str:
    if file_type == "txt" or file_type == "csv":
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    if file_type == "pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(path)
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            return ""
    if file_type == "docx":
        try:
            import docx
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    if file_type == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        except Exception:
            return ""
    return ""


def _chunk_text(text: str, size: int, overlap: int) -> List[str]:
    # Chunk by line, not by a flat word list — joining an entire document's
    # words with single spaces destroys row/line structure, which silently
    # corrupts tabular data (CSV/XLSX rows bleed into each other).
    lines = text.split("\n")
    chunks: List[str] = []
    current: List[str] = []
    count = 0
    for line in lines:
        line_words = len(line.split())
        if count + line_words > size and current:
            chunks.append("\n".join(current))
            tail: List[str] = []
            tail_count = 0
            for prev_line in reversed(current):
                tail_count += len(prev_line.split())
                tail.insert(0, prev_line)
                if tail_count >= overlap:
                    break
            current, count = tail, tail_count
        current.append(line)
        count += line_words
    if current:
        chunks.append("\n".join(current))
    return chunks


def _assert_public_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise HTTPException(status_code=400, detail="Only http/https URLs are supported")
    if not parsed.hostname:
        raise HTTPException(status_code=400, detail="Invalid URL")

    # Basic SSRF guard: a business owner could otherwise point this at
    # localhost/internal services/cloud metadata endpoints. Resolve the
    # hostname and reject anything that isn't a genuine public address.
    try:
        resolved_ip = socket.gethostbyname(parsed.hostname)
    except socket.gaierror:
        raise HTTPException(status_code=400, detail="Could not resolve URL host")

    ip = ipaddress.ip_address(resolved_ip)
    if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast:
        raise HTTPException(status_code=400, detail="URLs pointing to private/internal addresses are not allowed")


def _site_root(url: str) -> str:
    parsed = urlparse(url)
    return f"{parsed.scheme}://{parsed.netloc}"


async def _get_robot_parser(base_url: str) -> RobotFileParser:
    parser = RobotFileParser()
    robots_url = urljoin(base_url + "/", "robots.txt")
    try:
        async with httpx.AsyncClient(timeout=8) as client:
            resp = await client.get(robots_url, headers={"User-Agent": CRAWL_USER_AGENT})
        if resp.status_code >= 400:
            parser.parse([])  # no robots.txt -- allow everything
        else:
            parser.parse(resp.text.splitlines())
    except httpx.HTTPError:
        parser.parse([])  # unreachable robots.txt -- fail open, same as "not present"
    return parser


def _looks_like_page(url: str) -> bool:
    path = urlparse(url).path.lower()
    if any(path.endswith(ext) for ext in _NON_PAGE_EXTENSIONS):
        return False
    return True


async def _fetch_sitemap_xml(sitemap_url: str, _depth: int = 0) -> List[str]:
    """Fetches and parses one sitemap file at an already-complete URL --
    follows one level of <sitemapindex> nesting. Best-effort: returns []
    on any failure rather than raising, since a missing/broken sitemap
    just means falling back to a link crawl."""
    if _depth > 1:  # one level of sitemap-index nesting is enough
        return []
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.get(sitemap_url, headers={"User-Agent": CRAWL_USER_AGENT})
        if resp.status_code >= 400 or not resp.text.strip():
            return []
        root = ElementTree.fromstring(resp.content)
    except (httpx.HTTPError, ElementTree.ParseError):
        return []

    tag = root.tag.lower()
    if tag.endswith("sitemapindex"):
        nested = [el.text.strip() for el in root.iter() if el.tag.lower().endswith("loc") and el.text]
        urls: List[str] = []
        for nested_url in nested[:5]:  # bounded -- don't chase an unbounded index
            urls.extend(await _fetch_sitemap_xml(nested_url, _depth + 1))
            if len(urls) >= MAX_CRAWL_PAGES:
                break
        return urls
    return [el.text.strip() for el in root.iter() if el.tag.lower().endswith("loc") and el.text]


async def _discover_sitemap_urls(base_url: str) -> List[str]:
    """base_url is a site root (e.g. https://example.com) -- resolves it to
    /sitemap.xml once, then hands off to _fetch_sitemap_xml for the actual
    fetch+parse(+nested-index-following)."""
    sitemap_url = urljoin(base_url + "/", "sitemap.xml")
    return await _fetch_sitemap_xml(sitemap_url)


async def _discover_by_crawling(base_url: str, max_pages: int = MAX_CRAWL_PAGES) -> List[str]:
    from bs4 import BeautifulSoup

    domain = urlparse(base_url).netloc
    seen = {base_url}
    queue = [(base_url, 0)]
    found: List[str] = []

    async with httpx.AsyncClient(timeout=10) as client:
        while queue and len(found) < max_pages:
            url, depth = queue.pop(0)
            try:
                _assert_public_url(url)
                resp = await client.get(url, headers={"User-Agent": CRAWL_USER_AGENT})
            except (httpx.HTTPError, HTTPException):
                continue
            if resp.status_code >= 400 or "text/html" not in resp.headers.get("content-type", ""):
                continue
            found.append(url)
            if depth >= 2:
                continue
            soup = BeautifulSoup(resp.text, "html.parser")
            for a in soup.find_all("a", href=True):
                link = urljoin(url, a["href"]).split("#")[0]
                parsed = urlparse(link)
                if parsed.netloc != domain or link in seen or not _looks_like_page(link):
                    continue
                seen.add(link)
                if len(seen) <= max_pages * 4:  # bound queue growth on link-heavy pages
                    queue.append((link, depth + 1))
    return found


async def discover_website_pages(url: str) -> List[str]:
    """Sitemap first, link-crawl fallback; filtered by robots.txt and capped
    at MAX_CRAWL_PAGES either way."""
    _assert_public_url(url)
    base_url = _site_root(url)

    candidates = await _discover_sitemap_urls(base_url)
    if not candidates:
        candidates = await _discover_by_crawling(base_url)

    robots = await _get_robot_parser(base_url)
    seen = set()
    pages: List[str] = []
    for page_url in candidates:
        if page_url in seen or not _looks_like_page(page_url):
            continue
        seen.add(page_url)
        if urlparse(page_url).netloc != urlparse(base_url).netloc:
            continue
        if not robots.can_fetch(CRAWL_USER_AGENT, page_url):
            continue
        pages.append(page_url)
        if len(pages) >= MAX_CRAWL_PAGES:
            break
    return pages


async def _fetch_url_text(url: str) -> str:
    from bs4 import BeautifulSoup

    # Redirects are followed manually (never via httpx's follow_redirects)
    # and each hop is re-validated with _assert_public_url. Otherwise a
    # business-controlled public URL could 302 to a private/internal
    # address or cloud metadata endpoint and have that response ingested —
    # the one-time check on the original URL wouldn't catch that.
    max_redirects = 5
    current_url = url
    async with httpx.AsyncClient(timeout=15, follow_redirects=False) as client:
        for _ in range(max_redirects + 1):
            _assert_public_url(current_url)
            try:
                resp = await client.get(current_url, headers={"User-Agent": "MielikkiXBot/1.0"})
            except httpx.HTTPError:
                raise HTTPException(status_code=400, detail="Could not fetch that URL")
            if resp.is_redirect:
                location = resp.headers.get("location")
                if not location:
                    raise HTTPException(status_code=400, detail="Invalid redirect response")
                current_url = str(httpx.URL(current_url).join(location))
                continue
            break
        else:
            raise HTTPException(status_code=400, detail="Too many redirects")
    if resp.status_code >= 400:
        raise HTTPException(status_code=400, detail=f"URL returned status {resp.status_code}")
    if len(resp.content) > MAX_URL_FETCH_BYTES:
        raise HTTPException(status_code=400, detail="Page is too large")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    lines = [line.strip() for line in soup.get_text(separator="\n").splitlines() if line.strip()]
    return "\n".join(lines)


async def ingest_url(db: Session, business_id: str, user_id: str, url: str) -> Document:
    text = await _fetch_url_text(url)

    doc = Document(
        business_id=business_id,
        filename=url,
        file_url=url,
        file_type="url",
        status="processing",
        uploaded_by=user_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    _embed_and_store(db, doc, text)
    return doc


async def crawl_and_ingest_website(business_id: str, user_id: str, urls: List[str]) -> None:
    """Background-task worker for POST /api/documents/from-website. Runs
    after the response is already sent, so it opens its own DB session --
    the request's injected session is closed by then (see get_db's
    `finally: db.close()`). Each page reuses ingest_url unchanged; a single
    bad page (404, timeout, odd encoding) is skipped rather than aborting
    the rest of the batch."""
    db = SessionLocal()
    try:
        from ..models.business import Business

        business = db.query(Business).filter(Business.id == business_id).first()
        if not business:
            return

        for url in urls:
            plan = plan_service.get_plan(business.plan)
            if plan.limits.max_document_uploads is not None:
                current_count = db.query(Document).filter(Document.business_id == business_id).count()
                if current_count >= plan.limits.max_document_uploads:
                    break  # plan cap reached mid-crawl -- stop, don't raise (nothing to return this to)

            already_exists = (
                db.query(Document)
                .filter(Document.business_id == business_id, Document.filename == url)
                .first()
            )
            if already_exists:
                continue

            try:
                await ingest_url(db, business_id, user_id, url)
            except Exception:
                continue  # one bad page shouldn't abort the rest of the crawl

            await asyncio.sleep(0.5)  # politeness toward the target site, not a security control
    finally:
        db.close()


async def ingest_document(
    db: Session, business_id: str, user_id: str, file: UploadFile
) -> Document:
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail=f"File type .{ext} not supported")

    existing = (
        db.query(Document)
        .filter(Document.business_id == business_id, Document.filename == file.filename)
        .first()
    )
    if existing:
        raise HTTPException(
            status_code=400,
            detail=f"'{file.filename}' has already been uploaded. Delete the existing copy first if you want to replace it.",
        )

    content = await file.read()
    if len(content) > settings.max_upload_size_mb * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large")

    os.makedirs(settings.upload_dir, exist_ok=True)
    saved_name = f"{uuid.uuid4()}.{ext}"
    saved_path = os.path.join(settings.upload_dir, saved_name)
    with open(saved_path, "wb") as f:
        f.write(content)

    doc = Document(
        business_id=business_id,
        filename=file.filename,
        file_url=saved_path,
        file_type=ext,
        status="processing",
        uploaded_by=user_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    _process_document(db, doc, saved_path, ext)
    return doc


def _process_document(db: Session, doc: Document, path: str, ext: str):
    text = _extract_text(path, ext)
    _embed_and_store(db, doc, text)


def _embed_and_store(db: Session, doc: Document, text: str):
    try:
        chunks = _chunk_text(text, settings.chunk_size, settings.chunk_overlap)
        embeddings = embed_texts(chunks)

        for idx, (chunk_text, emb) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                business_id=doc.business_id,
                document_id=doc.id,
                chunk_index=idx,
                content=chunk_text,
                embedding_json=json.dumps(emb),
            )
            db.add(chunk)

        doc.status = "embedded"
    except Exception:
        doc.status = "failed"
    finally:
        db.commit()
